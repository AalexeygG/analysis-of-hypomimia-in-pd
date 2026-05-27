from __future__ import annotations
from typing import Optional
import csv
from pathlib import Path
from docx import Document
from docx.shared import Inches


def _fmt_stat(x):
    return "н/д" if x is None else x


def build_csv_report(
    csv_path: str,
    features: dict,
    stats: dict,
    test_type: str,
    group: str,
    quality_score: float,
    conclusion_text: str,
    risk_score: float,
):
    """
    CSV-отчёт: признаки + U/p-value + итоговая интерпретация.
    """
    Path(csv_path).parent.mkdir(parents=True, exist_ok=True)

    with open(csv_path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f, delimiter=";")
        w.writerow(["Проба", test_type])
        w.writerow(["Группа", group])
        w.writerow(["Качество трекинга (доля кадров с лицом)", f"{quality_score:.3f}"])
        w.writerow(["Итог", conclusion_text])
        w.writerow(["Оценка риска (0..1)", f"{risk_score:.3f}"])
        w.writerow([])

        w.writerow(["Признак", "Значение", "U (Манна—Уитни)", "p-value"])
        for k, v in features.items():
            u = stats.get(k, {}).get("U", None)
            p = stats.get(k, {}).get("p_value", None)
            w.writerow([k, f"{float(v):.6f}", _fmt_stat(u), _fmt_stat(p)])


def build_docx_report(
    docx_path: str,
    features: dict,
    stats: dict,
    test_type: str,
    group: str,
    quality_score: float,
    dynamics_plot_path: Optional[str] = None,
    box_plot_path: Optional[str] = None,
    conclusion_text: str = "",
    risk_score: float = 0.0,
):
    """
    DOCX-отчёт: таблицы + кликабельность на сайте обеспечивается отдельно.
    В DOCX вставляем оба графика как картинки.
    """
    Path(docx_path).parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("Отчёт обработки видеозаписи", level=1)

    # Общие сведения
    doc.add_paragraph(f"Проба: {test_type}")
    doc.add_paragraph(f"Группа: {group}")
    doc.add_paragraph(f"Качество трекинга (доля кадров с лицом): {quality_score:.3f}")

    # Итог
    if conclusion_text:
        doc.add_heading("Итоговая интерпретация", level=2)
        doc.add_paragraph(f"Итог: {conclusion_text}")
        doc.add_paragraph(f"Оценка риска (0…1): {risk_score:.3f}")

    # Таблица признаков + статистики
    doc.add_heading("Признаки и статистика", level=2)
    doc.add_paragraph("Вектор признаков и результаты статистического сравнения (Манна—Уитни):")

    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Признак"
    hdr[1].text = "Значение"
    hdr[2].text = "U"
    hdr[3].text = "p-value"

    any_nd = False
    for k, v in features.items():
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = f"{float(v):.6f}"

        u = stats.get(k, {}).get("U", None)
        p = stats.get(k, {}).get("p_value", None)

        if u is None or p is None:
            any_nd = True

        row[2].text = "н/д" if u is None else str(u)
        row[3].text = "н/д" if p is None else str(p)

    if any_nd:
        doc.add_paragraph(
            "Примечание: значения U и p-value обозначены как «н/д», если объём данных в одной из групп "
            "недостаточен для корректного расчёта (обычно требуется не менее двух наблюдений в каждой группе "
            "для выбранной пробы)."
        )

    # Графики
    if dynamics_plot_path and Path(dynamics_plot_path).exists():
        doc.add_heading("Динамика индекса (данная запись)", level=2)
        doc.add_picture(dynamics_plot_path, width=Inches(6.2))

    if box_plot_path and Path(box_plot_path).exists():
        doc.add_heading("Boxplot признаков (сравнение групп)", level=2)
        doc.add_picture(box_plot_path, width=Inches(6.2))

    doc.save(docx_path)
